#!/usr/bin/env python3
"""
Documentation Generator for Cross-Domain Predictive Analytics Dashboard

This script creates a comprehensive DOCX report documenting the project,
including screenshots, code samples, and architectural diagrams.

Usage:
    python generate_documentation.py

Requirements:
    - python-docx
    - Pillow (for image processing)
    - selenium (for capturing screenshots)
    - The application should be running on localhost:5000
"""

import os
import time
import glob
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
from PIL import Image
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# Configuration
APP_URL = "http://localhost:5000"
SCREENSHOT_DIR = "documentation/screenshots"
OUTPUT_FILE = "Cross_Domain_Predictive_Analytics_Dashboard_Documentation.docx"

# Team members
TEAM_MEMBERS = {
    "Ade": "System Integration & Real-Time Analytics",
    "Rujeko": "Frontend Development", 
    "Emmanuel": "Data Visualization",
    "Julie": "API Integration & Data Processing",
    "Chao": "Machine Learning & Predictive Modeling"
}

# Pages to capture
PAGES = [
    {"name": "Home", "url": "/", "wait_for": "navbar"},
    {"name": "Dashboard", "url": "/dashboard", "wait_for": "domainTabs"},
    {"name": "Natural Language Queries", "url": "/nlq", "wait_for": "query-form"},
    {"name": "System Status", "url": "/system/system-status", "wait_for": "status-panel"},
    {"name": "Demo Control", "url": "/demo/control", "wait_for": "demo-controls"}
]

def ensure_dir(directory):
    """Ensure that a directory exists."""
    if not os.path.exists(directory):
        os.makedirs(directory)

def take_screenshots():
    """Take screenshots of all pages in the application."""
    print("Taking screenshots...")
    ensure_dir(SCREENSHOT_DIR)
    
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--window-size=1920,1080")
    
    driver = webdriver.Chrome(options=chrome_options)
    
    screenshots = []
    
    try:
        for page in PAGES:
            url = f"{APP_URL}{page['url']}"
            print(f"Capturing {page['name']} at {url}")
            
            driver.get(url)
            
            # Wait for the page to load completely
            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.ID, page['wait_for']))
                )
                # Give an extra second for animations to complete
                time.sleep(1)
            except Exception as e:
                print(f"Warning: Timed out waiting for {page['wait_for']} element on {page['name']}: {e}")
            
            screenshot_path = os.path.join(SCREENSHOT_DIR, f"{page['name'].lower().replace(' ', '_')}.png")
            driver.save_screenshot(screenshot_path)
            screenshots.append({"name": page['name'], "path": screenshot_path})
            
            # If it's the dashboard, capture individual tabs
            if page['name'] == "Dashboard":
                # Get all tab links
                tab_links = driver.find_elements(By.CSS_SELECTOR, "#domainTabs .nav-link")
                
                for i, tab_link in enumerate(tab_links):
                    tab_name = tab_link.text.strip()
                    if tab_name:  # Skip empty tabs
                        try:
                            tab_link.click()
                            time.sleep(1)  # Wait for tab content to load
                            tab_screenshot_path = os.path.join(SCREENSHOT_DIR, f"dashboard_tab_{tab_name.lower().replace(' ', '_')}.png")
                            driver.save_screenshot(tab_screenshot_path)
                            screenshots.append({"name": f"Dashboard - {tab_name} Tab", "path": tab_screenshot_path})
                        except Exception as e:
                            print(f"Error capturing tab {tab_name}: {e}")
    
    finally:
        driver.quit()
        
    return screenshots

def generate_docx(screenshots):
    """Generate a comprehensive DOCX document with documentation and screenshots."""
    print("Generating documentation...")
    
    doc = Document()
    
    # Set up document properties
    core_properties = doc.core_properties
    core_properties.author = "Cross-Domain Predictive Analytics Team"
    core_properties.title = "Cross-Domain Predictive Analytics Dashboard Documentation"
    core_properties.comments = "Comprehensive documentation of the project"
    
    # Title Page
    title = doc.add_heading('Cross-Domain Predictive Analytics Dashboard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Comprehensive Project Documentation')
    subtitle_run.font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f'Generated on {datetime.datetime.now().strftime("%B %d, %Y")}')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add team members
    doc.add_paragraph()
    team_heading = doc.add_heading('Project Team', level=1)
    team_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    team_table = doc.add_table(rows=1, cols=2)
    team_table.style = 'Table Grid'
    team_table.cell(0, 0).text = "Team Member"
    team_table.cell(0, 1).text = "Role"
    
    for member, role in TEAM_MEMBERS.items():
        row = team_table.add_row()
        row.cells[0].text = member
        row.cells[1].text = role
    
    # Add page break
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_entries = [
        "1. Executive Summary",
        "2. Project Overview",
        "3. System Architecture",
        "4. Individual Components",
        "   4.1. Frontend Development (Rujeko)",
        "   4.2. Data Visualization (Emmanuel)",
        "   4.3. API Integration (Julie)",
        "   4.4. Machine Learning Models (Chao)",
        "   4.5. System Integration (Ade)",
        "5. User Guide",
        "6. Implementation Details",
        "7. Testing and Validation",
        "8. Conclusions and Future Work"
    ]
    
    for entry in toc_entries:
        p = doc.add_paragraph()
        p.add_run(entry)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph('The Cross-Domain Predictive Analytics Dashboard is a sophisticated web-based platform that integrates data from multiple domains—weather forecasts, economic indicators, social media trends, and transportation metrics—to provide predictive analytics and actionable insights. The system uses machine learning algorithms to identify patterns across seemingly unrelated datasets, enabling users to make proactive decisions based on correlated data insights.')
    doc.add_paragraph('This project demonstrates how interdisciplinary data integration can reveal non-obvious relationships between different domains, allowing organizations to anticipate changes and optimize their operations accordingly. The system supports various use cases including supply chain optimization, public health response planning, urban infrastructure management, and financial market strategy.')
    
    # Project Overview
    doc.add_heading('2. Project Overview', level=1)
    doc.add_paragraph("The project involved developing an advanced web-based data platform using Python's Flask framework. The dashboard integrates multiple data sources from public APIs and applies machine learning models to deliver predictive analytics and actionable insights. The focus was on cross-domain data correlation, predictive modeling, API integration, and interactive visualization.")
    
    p = doc.add_paragraph('Key features include:')
    doc.add_paragraph('• Multi-API integration with data sources across different domains', style='List Bullet')
    doc.add_paragraph('• Machine learning models for prediction and cross-domain correlation', style='List Bullet')
    doc.add_paragraph('• Interactive data visualization with confidence indicators', style='List Bullet')
    doc.add_paragraph('• Natural language query capabilities', style='List Bullet')
    doc.add_paragraph('• Real-time updates and alerting system', style='List Bullet')
    doc.add_paragraph('• Cross-domain correlation analysis', style='List Bullet')
    
    # System Architecture
    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph('The system follows a modular architecture with clear separation of concerns:')
    
    architecture_table = doc.add_table(rows=1, cols=2)
    architecture_table.style = 'Table Grid'
    architecture_table.cell(0, 0).text = "Component"
    architecture_table.cell(0, 1).text = "Description"
    
    components = [
        ("Frontend Layer", "User interface built with Flask templates, HTML, CSS, and JavaScript. Features responsive design and interactive components."),
        ("API Integration Layer", "Connects to external APIs with rate limiting, caching, and error handling."),
        ("Data Processing Layer", "Cleans, transforms, and correlates data from different sources."),
        ("Machine Learning Layer", "Implements predictive models and cross-domain correlation analysis."),
        ("Visualization Layer", "Renders data using interactive charts, graphs, and maps."),
        ("System Integration Layer", "Coordinates all components and manages real-time updates via WebSockets.")
    ]
    
    for component, description in components:
        row = architecture_table.add_row()
        row.cells[0].text = component
        row.cells[1].text = description
    
    # Add screenshots after introducing the architecture
    doc.add_heading('Application Screenshots', level=2)
    
    for screenshot in screenshots:
        doc.add_heading(screenshot["name"], level=3)
        try:
            doc.add_picture(screenshot["path"], width=Inches(6))
            paragraph = doc.paragraphs[-1]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'Figure: {screenshot["name"]} view')
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.style.font.italic = True
        except Exception as e:
            doc.add_paragraph(f"Error adding screenshot {screenshot['path']}: {e}")
    
    # Individual Components
    doc.add_heading('4. Individual Components', level=1)
    
    # 4.1 Frontend Development (Rujeko)
    doc.add_heading('4.1. Frontend Development (Rujeko)', level=2)
    doc.add_paragraph('The frontend was developed using HTML, CSS, and JavaScript with the Bootstrap framework for responsive design. Key features include:')
    doc.add_paragraph('• Responsive layout that works well on different devices', style='List Bullet')
    doc.add_paragraph('• Interactive controls for data filtering and exploration', style='List Bullet')
    doc.add_paragraph('• Tab-based interface for navigating between domains', style='List Bullet')
    doc.add_paragraph('• Real-time updates via WebSockets', style='List Bullet')
    doc.add_paragraph('• Natural language query input with instant results', style='List Bullet')
    
    # 4.2 Data Visualization (Emmanuel)
    doc.add_heading('4.2. Data Visualization (Emmanuel)', level=2)
    doc.add_paragraph('The visualization components were implemented using libraries like Plotly and D3.js, providing interactive and informative visual representations of data:')
    doc.add_paragraph('• Time series charts for historical and predicted data', style='List Bullet')
    doc.add_paragraph('• Correlation heatmaps and network diagrams', style='List Bullet')
    doc.add_paragraph('• Geospatial visualizations for location-based data', style='List Bullet')
    doc.add_paragraph('• Confidence interval visualizations', style='List Bullet')
    doc.add_paragraph('• Dashboard formatters for different types of data', style='List Bullet')
    
    # 4.3 API Integration (Julie)
    doc.add_heading('4.3. API Integration (Julie)', level=2)
    doc.add_paragraph('The API integration layer connects to multiple external data sources while handling rate limiting, caching, and error recovery:')
    doc.add_paragraph('• Base connector system with common functionality', style='List Bullet')
    doc.add_paragraph('• Domain-specific connectors for weather, economic, transportation, and social media data', style='List Bullet')
    doc.add_paragraph('• Caching mechanism to reduce API calls', style='List Bullet')
    doc.add_paragraph('• Error handling and fallback mechanisms', style='List Bullet')
    doc.add_paragraph('• Rate limiting to prevent API quota exhaustion', style='List Bullet')
    
    # 4.4 Machine Learning Models (Chao)
    doc.add_heading('4.4. Machine Learning Models (Chao)', level=2)
    doc.add_paragraph('The machine learning components implement predictive models and cross-domain correlation analysis:')
    doc.add_paragraph('• LSTM models for time series prediction', style='List Bullet')
    doc.add_paragraph('• Cross-domain correlation algorithms', style='List Bullet')
    doc.add_paragraph('• Confidence scoring for prediction reliability', style='List Bullet')
    doc.add_paragraph('• "What-if" scenario modeling', style='List Bullet')
    doc.add_paragraph('• Model training pipelines with data preprocessing', style='List Bullet')
    
    # 4.5 System Integration (Ade)
    doc.add_heading('4.5. System Integration (Ade)', level=2)
    doc.add_paragraph('The system integration layer coordinates all components and ensures smooth operation:')
    doc.add_paragraph('• Real-time data pipeline with WebSocket updates', style='List Bullet')
    doc.add_paragraph('• Event-based architecture for component communication', style='List Bullet')
    doc.add_paragraph('• Alert system for prediction thresholds', style='List Bullet')
    doc.add_paragraph('• System status monitoring', style='List Bullet')
    doc.add_paragraph('• Component registration and discovery', style='List Bullet')
    
    # User Guide
    doc.add_heading('5. User Guide', level=1)
    doc.add_paragraph('This section provides detailed instructions for using the Cross-Domain Predictive Analytics Dashboard:')
    
    # Navigation
    doc.add_heading('Dashboard Navigation', level=2)
    doc.add_paragraph('The main dashboard provides a comprehensive view of all domains with the following features:')
    doc.add_paragraph('• Overview tab: Summary of all domains with key metrics', style='List Bullet')
    doc.add_paragraph('• Domain-specific tabs: Detailed view of each domain's data', style='List Bullet')
    doc.add_paragraph('• Cross-Domain tab: Correlation analysis across domains', style='List Bullet')
    doc.add_paragraph('• Time range selector: Filter data by time period', style='List Bullet')
    doc.add_paragraph('• Auto-refresh toggle: Enable/disable automatic updates', style='List Bullet')
    
    # Natural Language Queries
    doc.add_heading('Natural Language Queries', level=2)
    doc.add_paragraph('The system supports natural language queries, allowing users to ask questions in plain English:')
    doc.add_paragraph('• Example: "Show me the correlation between temperature and traffic congestion"', style='List Bullet')
    doc.add_paragraph('• Example: "Predict economic indicators for next week based on weather forecasts"', style='List Bullet')
    doc.add_paragraph('• Example: "What would happen to traffic if temperature increases by 10 degrees?"', style='List Bullet')
    
    # Implementation Details
    doc.add_heading('6. Implementation Details', level=1)
    doc.add_paragraph('The dashboard is implemented using a modern stack of technologies:')
    doc.add_paragraph('• Backend: Python Flask framework with SQLite for persistence', style='List Bullet')
    doc.add_paragraph('• Frontend: HTML, CSS, JavaScript with Bootstrap for responsive design', style='List Bullet')
    doc.add_paragraph('• Real-time updates: Flask-SocketIO for WebSocket communication', style='List Bullet')
    doc.add_paragraph('• Data processing: Pandas and NumPy for data manipulation', style='List Bullet')
    doc.add_paragraph('• Machine learning: TensorFlow/Keras for LSTM models, scikit-learn for analysis', style='List Bullet')
    doc.add_paragraph('• Visualization: Plotly and D3.js for interactive charts', style='List Bullet')
    
    # Testing and Validation
    doc.add_heading('7. Testing and Validation', level=1)
    doc.add_paragraph('The system underwent comprehensive testing to ensure reliability and accuracy:')
    doc.add_paragraph('• Unit tests for individual components', style='List Bullet')
    doc.add_paragraph('• Integration tests for component interactions', style='List Bullet')
    doc.add_paragraph('• Machine learning model validation using historical data', style='List Bullet')
    doc.add_paragraph('• Browser testing for frontend compatibility', style='List Bullet')
    doc.add_paragraph('• Performance testing under various loads', style='List Bullet')
    
    # Conclusions and Future Work
    doc.add_heading('8. Conclusions and Future Work', level=1)
    doc.add_paragraph('The Cross-Domain Predictive Analytics Dashboard successfully demonstrates the value of integrating data from multiple domains to enable predictive analytics and actionable insights. By identifying patterns across seemingly unrelated datasets, the system provides users with valuable foresight for proactive decision-making.')
    
    doc.add_paragraph('Future work could include:')
    doc.add_paragraph('• Integration with additional data sources and domains', style='List Bullet')
    doc.add_paragraph('• More advanced machine learning models for improved prediction accuracy', style='List Bullet')
    doc.add_paragraph('• Enhanced natural language processing for more complex queries', style='List Bullet')
    doc.add_paragraph('• Mobile application for on-the-go access', style='List Bullet')
    doc.add_paragraph('• Integration with notification systems (email, SMS, etc.)', style='List Bullet')
    
    # Save the document
    doc.save(OUTPUT_FILE)
    print(f"Documentation saved to {OUTPUT_FILE}")

def main():
    """Main function to generate documentation."""
    print("Starting documentation generation...")
    ensure_dir("documentation")
    
    try:
        screenshots = take_screenshots()
        generate_docx(screenshots)
        print("Documentation generation complete!")
    except Exception as e:
        print(f"Error generating documentation: {e}")

if __name__ == "__main__":
    main() 